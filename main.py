import streamlit as st
import google.generativeai as genai
import pyttsx3
from gtts import gTTS
import pytesseract
from PIL import Image
import pdfplumber
import docx2txt
import io
import requests
import threading
import random
from googletrans import Translator
from docx import Document
from fpdf import FPDF
import os

# üîπ API Keys Rotation to Avoid Quota Limits
API_KEYS = [
    "AIzaSyCh3M8GBfAmbc3p1ySkAAT7jbI5LrBy8NI",
    "AIzaSyBGykxdPER4wAIkM_GJbpBR_nY2yv5_zVU",
    "AIzaSyApqTuxNCW_TtW_4O28ZYtXDYGV--MBlCw"
]


def get_random_api_key():
    return random.choice(API_KEYS)


# üîπ Configure Tesseract OCR Path
pytesseract.pytesseract.tesseract_cmd = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"

# üîπ Google Translate Fallback
translator = Translator()

# üîπ Language Code Mapping for gTTS
LANGUAGE_CODES = {
    "English": "en", "Spanish": "es", "French": "fr", "German": "de", "Hindi": "hi",
    "Chinese (Simplified)": "zh-cn", "Japanese": "ja", "Portuguese": "pt", "Russian": "ru", "Arabic": "ar",
    "Italian": "it", "Korean": "ko", "Dutch": "nl", "Swedish": "sv", "Turkish": "tr", "Persian (Iran)": "fa",
    "Punjabi": "pa"
}

COUNTRY_LANGUAGES = {
    "United States": ("English", "Spanish"), "India": ("Hindi", "Punjabi"), "France": ("French", "English"),
    "Germany": ("German", "English"), "China": ("Chinese (Simplified)", "English"), "Japan": ("Japanese", "English"),
    "Brazil": ("Portuguese", "English"), "Russia": ("Russian", "English"), "Saudi Arabia": ("Arabic", "English"),
    "Spain": ("Spanish", "English"), "Italy": ("Italian", "English"), "South Korea": ("Korean", "English"),
    "Pakistan": ("Urdu", "English"), "Netherlands": ("Dutch", "English"), "Sweden": ("Swedish", "English"),
    "Turkey": ("Turkish", "English"), "Iran": ("Persian (Iran)", "English")
}


def speak_text(text, language):
    try:
        lang_code = LANGUAGE_CODES.get(language, "en")  # Default to English if language not found
        tts = gTTS(text=text, lang=lang_code)
        tts.save("temp.mp3")
        st.audio("temp.mp3", format="audio/mp3")
    except Exception as e:
        st.warning(f"Error in speech synthesis: {str(e)}")


# üîπ Extract Text from Documents

def extract_text_from_file(uploaded_file):
    text = ""
    if uploaded_file.name.endswith(".txt"):
        text = uploaded_file.getvalue().decode("utf-8")
    elif uploaded_file.name.endswith(".pdf"):
        with pdfplumber.open(uploaded_file) as pdf:
            text = "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())
    elif uploaded_file.name.endswith(".docx"):
        text = docx2txt.process(uploaded_file)
    return text.strip() if text else "No text found in the file."

# üîπ Extract Text from Image Using OCR
def extract_text_from_image(uploaded_image):
    image = Image.open(uploaded_image)
    text = pytesseract.image_to_string(image)
    return text.strip() if text else "No text found in the image."

# üîπ Generate Downloadable File
def generate_downloadable_file(content, original_filename):
    file_extension = original_filename.split(".")[-1]
    buffer = io.BytesIO()
    if file_extension == "txt":
        buffer.write(content.encode("utf-8"))
    elif file_extension == "docx":
        doc = Document()
        doc.add_paragraph(content)
        doc.save(buffer)
    elif file_extension == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.add_font("DejaVu", "", "DejaVuSans.ttf", uni=True)
        pdf.set_font("DejaVu", size=12)
        pdf.multi_cell(0, 10, content)
        pdf.output(buffer)
    buffer.seek(0)
    return buffer, f"translated.{file_extension}", f"application/{file_extension}"

def ensure_font():
    font_path = os.path.abspath("DejaVuSans.ttf")
    if not os.path.exists(font_path):
        url = "https://github.com/dejavu-fonts/dejavu-fonts/raw/master/ttf/DejaVuSans.ttf"
        response = requests.get(url)
        with open(font_path, "wb") as f:
            f.write(response.content)
    return font_path  # Return absolute path

# üîπ Generate Downloadable File
# üîπ Generate Downloadable File
def generate_downloadable_file(content, original_filename):
    file_extension = original_filename.split(".")[-1]
    buffer = io.BytesIO()

    if file_extension == "txt":
        buffer.write(content.encode("utf-8"))
    elif file_extension == "docx":
        doc = Document()
        doc.add_paragraph(content)
        doc.save(buffer)
    elif file_extension == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)  # Using built-in font instead of external one
        pdf.multi_cell(0, 10, content)
        pdf.output(buffer, "F")  # Ensure buffer contains valid PDF data

    buffer.seek(0)
    return buffer, f"translated.{file_extension}", f"application/{file_extension}"




# üîπ Streamlit App
def main():
    st.set_page_config(page_title="Multi-Mode Language Translator", page_icon="üó£Ô∏è")
    st.title("üåç SpeakEase AI - Multi-Mode Translator")

    ensure_font()

    # Sidebar for Language Selection
    st.sidebar.header("üåè Language Settings")
    all_languages = list(LANGUAGE_CODES.keys())
    input_language = st.sidebar.selectbox("Select Input Language:", all_languages)
    output_language = st.sidebar.selectbox("Select Output Language:", all_languages)

    tab1, tab2, tab3 = st.tabs(["üìÑ File Upload", "üñºÔ∏è Image Upload", "‚úçÔ∏è Manual Text"])

    with tab1:
        st.subheader("üìÑ Upload a Document")
        uploaded_file = st.file_uploader("Upload a document (PDF, DOCX, TXT):", type=["pdf", "docx", "txt"])
        if uploaded_file:
            extracted_text = extract_text_from_file(uploaded_file)
            translated_text = translator.translate(extracted_text, src=LANGUAGE_CODES.get(input_language, "en"),
                                                   dest=LANGUAGE_CODES.get(output_language, "en")).text
            st.text_area("Translated Text:", translated_text, height=150)
            speak_text(translated_text, output_language)

            buffer, filename, mime_type = generate_downloadable_file(translated_text, uploaded_file.name)
            st.download_button("Download Translated File", buffer, file_name=filename, mime=mime_type,
                               key="download_file")

    with tab2:
        st.subheader("üñºÔ∏è Upload an Image")
        uploaded_image = st.file_uploader("Upload an image (PNG, JPG, JPEG):", type=["png", "jpg", "jpeg"])
        if uploaded_image:
            extracted_text = extract_text_from_image(uploaded_image)
            translated_text = translator.translate(extracted_text, src=LANGUAGE_CODES.get(input_language, "en"),
                                                   dest=LANGUAGE_CODES.get(output_language, "en")).text
            st.text_area("Translated Text:", translated_text, height=150)
            speak_text(translated_text, output_language)

            buffer, filename, mime_type = generate_downloadable_file(translated_text, "translated_image.txt")
            st.download_button("Download Translated Text", buffer, file_name=filename, mime=mime_type,
                               key="download_image")

    with tab3:
        st.subheader("‚úçÔ∏è Manual Text Translation")
        manual_text = st.text_area("Enter text to translate:", height=150)
        if st.button("Translate Text", key="translate_button"):
            translated_text = translator.translate(manual_text, src=LANGUAGE_CODES.get(input_language, "en"),
                                                   dest=LANGUAGE_CODES.get(output_language, "en")).text
            st.text_area("Translated Text:", translated_text, height=150)
            speak_text(translated_text, output_language)

            buffer, filename, mime_type = generate_downloadable_file(translated_text, "translated_manual.txt")
            st.download_button("Download Translated Text", buffer, file_name=filename, mime=mime_type,
                               key="download_manual")


# üöÄ Run Streamlit App
if __name__ == "__main__":
    main()
